from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches

from classcorpus.citations import format_citation
from classcorpus.database import Database
from classcorpus.indexer import sync_course
from classcorpus.parser_registry import ParserPlugin, ParserRegistry
from classcorpus.parsers import parse_source, supported_suffixes
from classcorpus.search import search
from tests.fixtures.make_fixtures import _png_bytes


def test_registry_rejects_suffix_conflicts_and_invalid_plugins():
    registry = ParserRegistry()
    plugin = ParserPlugin("first", (".one",), lambda path, output: [])
    registry.register(plugin)

    with pytest.raises(ValueError, match="already registered"):
        registry.register(ParserPlugin("second", (".ONE",), lambda path, output: []))
    with pytest.raises(ValueError, match="start with a dot"):
        registry.register(ParserPlugin("invalid", ("txt",), lambda path, output: []))

    assert registry.parser_for(".ONE") == plugin
    assert registry.supported_suffixes() == frozenset({".one"})


def test_markdown_plugin_preserves_raw_text_and_extracts_heading(tmp_path: Path):
    source = tmp_path / "lecture.md"
    raw_text = "# Greedy Algorithms\n\nExchange arguments preserve optimality.\n"
    source.write_text(raw_text, encoding="utf-8")

    record = parse_source(source, tmp_path / "unused")[0]

    assert record.title == "Greedy Algorithms"
    assert record.body_text == "Exchange arguments preserve optimality."
    assert record.raw_text == raw_text
    assert record.native_text_chars == len(raw_text)
    assert record.extraction_status == "text-extracted"
    assert record.render_path is None


def test_blank_text_document_is_explicitly_marked_for_review(tmp_path: Path):
    source = tmp_path / "blank.txt"
    source.write_text("\n  \n", encoding="utf-8")

    record = parse_source(source, tmp_path / "unused")[0]

    assert record.extraction_status == "review-needed"
    assert record.extraction_reasons == ("no-native-text",)


def test_docx_plugin_extracts_paragraphs_tables_and_image_risk(tmp_path: Path):
    source = tmp_path / "exam-practice.docx"
    document = Document()
    document.core_properties.title = "Calculus Practice"
    document.add_paragraph("Differentiate polynomial functions.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Function"
    table.cell(0, 1).text = "Derivative"
    table.cell(1, 0).text = "x squared"
    table.cell(1, 1).text = "2x"
    document.add_picture(BytesIO(_png_bytes()), width=Inches(1))
    document.save(source)

    record = parse_source(source, tmp_path / "unused")[0]

    assert record.title == "Calculus Practice"
    assert "Differentiate polynomial functions." in record.body_text
    assert "Function | Derivative" in record.body_text
    assert "x squared" in record.raw_text
    assert record.native_text_chars == len(record.raw_text)
    assert record.extraction_status == "review-needed"
    assert "embedded-image" in record.extraction_reasons
    assert record.has_visual_content is True
    assert record.render_path is None


def test_blank_docx_is_explicitly_marked_for_review(tmp_path: Path):
    source = tmp_path / "blank.docx"
    Document().save(source)

    record = parse_source(source, tmp_path / "unused")[0]

    assert record.extraction_status == "review-needed"
    assert record.extraction_reasons == ("no-native-text",)


def test_indexer_discovers_text_plugins_and_searches_with_citations(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    monkeypatch.setenv("CLASSCORPUS_DATA_DIR", str(tmp_path / "data"))
    root = tmp_path / "Algorithms"
    root.mkdir()
    (root / "lecture.md").write_text(
        "# Greedy Algorithms\nExchange argument matroid.",
        encoding="utf-8",
    )
    (root / "notes.txt").write_text(
        "Dynamic Programming\nOptimal substructure.",
        encoding="utf-8",
    )
    word_document = Document()
    word_document.add_heading("Exam Practice", level=1)
    word_document.add_paragraph("Differentiate with the polynomial power rule.")
    word_document.save(root / "practice.docx")
    (root / "ignored.csv").write_text("not,indexed", encoding="utf-8")
    database = Database(tmp_path / "index.sqlite3")
    database.initialize()

    report = sync_course(database, "Algorithms", root)
    greedy = search(database, "matroid", course="Algorithms")[0]
    dynamic = search(database, "optimal substructure", course="Algorithms")[0]
    practice = search(database, "polynomial power rule", course="Algorithms")[0]

    assert {".docx", ".md", ".txt"}.issubset(supported_suffixes())
    assert report.indexed == 3
    assert report.records_indexed == 3
    assert greedy.source_file == "lecture.md"
    assert greedy.ordinal == 1
    assert dynamic.source_file == "notes.txt"
    assert dynamic.ordinal == 1
    assert practice.source_file == "practice.docx"
    assert practice.ordinal == 1
    assert format_citation(practice) == "[Algorithms, practice.docx, Page 1]"
